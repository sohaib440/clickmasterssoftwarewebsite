"""Generate a Word document of full main-service page content."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Software-Mobile-Web-Development-Pages.docx"
SITE = "https://nextsoftwaredevelopment.com"
TARGET_SLUGS = (
    "software-development",
    "mobile-development",
    "web-development",
)

NAVY = RGBColor(0x0B, 0x1D, 0x36)
GOLD = RGBColor(0xB8, 0x94, 0x1F)
MUTED = RGBColor(0x47, 0x55, 0x69)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "0B1D36"
GOLD_HEX = "B8941F"
ROW_ALT = "F8F5EE"
CREAM = "F7F5F0"


def unescape(s: str) -> str:
    return (
        s.replace(r"\"", '"')
        .replace(r"\n", "\n")
        .replace(r"\\", "\\")
        .replace("&", "&")
        .strip()
    )


def slice_balanced(text: str, open_idx: int, open_ch: str = "{", close_ch: str = "}") -> str:
    depth = 0
    in_str = False
    quote = ""
    escape = False
    for i, ch in enumerate(text[open_idx:], open_idx):
        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == quote:
                in_str = False
            continue
        if ch in "\"'`":
            in_str = True
            quote = ch
            continue
        if ch == open_ch:
            depth += 1
        elif ch == close_ch:
            depth -= 1
            if depth == 0:
                return text[open_idx : i + 1]
    raise ValueError("Unbalanced block")


def extract_str(block: str, key: str) -> str:
    pattern = rf'{re.escape(key)}\s*:\s*"((?:[^"\\]|\\.)*)"'
    match = re.search(pattern, block)
    if match:
        return unescape(match.group(1))
    pattern = rf"{re.escape(key)}\s*:\s*`((?:[^`\\]|\\.)*)`"
    match = re.search(pattern, block)
    return unescape(match.group(1)) if match else ""


def extract_str_list(block: str, key: str) -> list[str]:
    match = re.search(rf"{re.escape(key)}\s*:\s*\[", block)
    if not match:
        return []
    arr = slice_balanced(block, match.end() - 1, "[", "]")
    return [unescape(item) for item in re.findall(r'"((?:[^"\\]|\\.)*)"', arr)]


def extract_objects(block: str, key: str) -> list[str]:
    match = re.search(rf"{re.escape(key)}\s*:\s*\[", block)
    if not match:
        return []
    arr = slice_balanced(block, match.end() - 1, "[", "]")
    objects: list[str] = []
    i = 0
    while True:
        start = arr.find("{", i)
        if start < 0:
            break
        obj = slice_balanced(arr, start)
        objects.append(obj)
        i = start + len(obj)
    return objects


def parse_define_categories(text: str) -> list[dict]:
    categories = []
    for match in re.finditer(r"defineCategory\(\{", text):
        block = slice_balanced(text, match.end() - 1)
        categories.append(
            {
                "slug": extract_str(block, "slug"),
                "label": extract_str(block, "label"),
                "exploreCta": extract_str(block, "exploreCta"),
                "pageTitle": extract_str(block, "pageTitle"),
                "tagline": extract_str(block, "tagline"),
                "description": extract_str(block, "description"),
                "metaDescription": extract_str(block, "metaDescription"),
                "heroImage": extract_str(block, "src"),
                "heroAlt": extract_str(block, "alt"),
                "highlights": [
                    {
                        "title": extract_str(obj, "title"),
                        "description": extract_str(obj, "description"),
                    }
                    for obj in extract_objects(block, "highlights")
                ],
                "faqs": [
                    {
                        "tag": extract_str(obj, "tag"),
                        "question": extract_str(obj, "question"),
                        "answer": extract_str(obj, "answer"),
                    }
                    for obj in extract_objects(block, "faqs")
                ],
            }
        )
    return categories


def parse_sub_services(text: str, slug: str) -> list[dict]:
    raw = re.search(r"const subServicesByCategoryRaw[^=]*=\s*\{", text)
    if not raw:
        raise SystemExit("Could not find subServicesByCategoryRaw")
    block = slice_balanced(text, raw.end() - 1)
    match = re.search(rf'"{re.escape(slug)}"\s*:\s*\[', block)
    if not match:
        raise SystemExit(f"Could not find {slug} sub-services")
    arr = slice_balanced(block, match.end() - 1, "[", "]")
    items = []
    i = 0
    while True:
        start = arr.find("sub({", i)
        if start < 0:
            break
        obj_start = arr.find("{", start)
        obj = slice_balanced(arr, obj_start)
        items.append(
            {
                "slug": extract_str(obj, "slug"),
                "label": extract_str(obj, "label"),
                "description": extract_str(obj, "description"),
                "tagline": extract_str(obj, "tagline"),
                "pageTitle": extract_str(obj, "pageTitle"),
                "metaDescription": extract_str(obj, "metaDescription"),
                "content": extract_str_list(obj, "content"),
            }
        )
        i = obj_start + len(obj)
    return items


def parse_named_object_array(text: str, name: str, keys: list[str]) -> list[dict]:
    match = re.search(rf"export const {re.escape(name)}\s*(?::[^=]+)?=\s*\[", text)
    if not match:
        return []
    arr = slice_balanced(text, match.end() - 1, "[", "]")
    items = []
    i = 0
    while True:
        start = arr.find("{", i)
        if start < 0:
            break
        obj = slice_balanced(arr, start)
        items.append({key: extract_str(obj, key) for key in keys})
        i = start + len(obj)
    return items


def parse_export_string(text: str, name: str) -> str:
    match = re.search(
        rf"export const {re.escape(name)}\s*=\s*\"((?:[^\"\\]|\\.)*)\"",
        text,
        re.DOTALL,
    )
    if match:
        return unescape(match.group(1))
    match = re.search(
        rf"export const {re.escape(name)}\s*=\s*\"((?:[^\"\\]|\\.)*)\"",
        text,
    )
    return unescape(match.group(1)) if match else ""


def parse_home_contact(text: str) -> dict:
    match = re.search(r"export const homeContact\s*=\s*\{", text)
    if not match:
        return {}
    block = slice_balanced(text, match.end() - 1)
    return {
        "overline": extract_str(block, "overline"),
        "titleBefore": extract_str(block, "titleBefore"),
        "titleEmphasis": extract_str(block, "titleEmphasis"),
        "subtext": extract_str(block, "subtext"),
        "cta": extract_str(block, "cta"),
        "reassurance": extract_str(block, "reassurance"),
    }


def parse_about(text: str) -> dict:
    match = re.search(r"export const aboutSection\s*=\s*\{", text)
    block = slice_balanced(text, match.end() - 1)
    values = []
    for obj in extract_objects(block, "values"):
        values.append(
            {
                "title": extract_str(obj, "title"),
                "description": extract_str(obj, "description"),
            }
        )
    paragraphs = extract_str_list(block, "paragraphs")
    return {
        "paragraphs": paragraphs,
        "values": values,
        "teamCta": extract_str(block, "teamCta"),
        "imageAlt": extract_str(block, "alt"),
    }


def parse_highlights(text: str) -> list[dict]:
    match = re.search(r"const defaultHighlights[^=]*=\s*\[", text)
    arr = slice_balanced(text, match.end() - 1, "[", "]")
    items = []
    i = 0
    while True:
        start = arr.find("{", i)
        if start < 0:
            break
        obj = slice_balanced(arr, start)
        items.append(
            {
                "title": extract_str(obj, "title"),
                "description": extract_str(obj, "description"),
            }
        )
        i = start + len(obj)
    return items


def parse_tech_groups(text: str) -> list[dict]:
    match = re.search(r"export const techStackLogoGroupMeta[^=]*=\s*\[", text)
    arr = slice_balanced(text, match.end() - 1, "[", "]")
    logos = {
        item["id"]: item["name"]
        for item in parse_named_object_array(text, "techStackLogos", ["id", "name"])
    }
    groups = []
    i = 0
    while True:
        start = arr.find("{", i)
        if start < 0:
            break
        obj = slice_balanced(arr, start)
        ids = extract_str_list(obj, "logoIds")
        groups.append(
            {
                "label": extract_str(obj, "label"),
                "description": extract_str(obj, "description"),
                "logos": [logos.get(logo_id, logo_id) for logo_id in ids],
            }
        )
        i = start + len(obj)
    return groups


def parse_project_highlights(text: str) -> dict[str, list[str]]:
    result: dict[str, list[str]] = {}
    for match in re.finditer(r'base\("([^"]+)"\)', text):
        slug = match.group(1)
        window = text[match.start() : match.start() + 8000]
        result[slug] = extract_str_list(window, "highlights")
    return result


def parse_project_slides(text: str) -> dict[str, list[dict]]:
    arrays: dict[str, list[dict]] = {}
    for match in re.finditer(r"const (\w+)\s*:\s*ProjectSlide\[\]\s*=\s*\[", text):
        name = match.group(1)
        arr = slice_balanced(text, match.end() - 1, "[", "]")
        slides = []
        i = 0
        while True:
            start = arr.find("{", i)
            if start < 0:
                break
            obj = slice_balanced(arr, start)
            img_match = re.search(
                r'img\(\s*"((?:[^"\\]|\\.)*)"\s*,\s*"((?:[^"\\]|\\.)*)"',
                obj,
            )
            slides.append(
                {
                    "label": extract_str(obj, "label"),
                    "caption": extract_str(obj, "caption"),
                    "src": unescape(img_match.group(1)) if img_match else "",
                    "alt": unescape(img_match.group(2)) if img_match else extract_str(obj, "alt"),
                }
            )
            i = start + len(obj)
        arrays[name] = slides

    result: dict[str, list[dict]] = {}
    for match in re.finditer(r"slides:\s*(\w+)", text):
        var_name = match.group(1)
        if var_name not in arrays:
            continue
        window = text[max(0, match.start() - 1200) : match.start()]
        bases = re.findall(r'base\("([^"]+)"\)', window)
        if bases:
            result[bases[-1]] = arrays[var_name]
    return result


def parse_case_study_cards(case_text: str, project_slugs: list[str]) -> list[dict]:
    match = re.search(r"export const projectCaseStudyMeta(?:\s*:[^=]+)?\s*=\s*\{", case_text)
    meta_block = slice_balanced(case_text, match.end() - 1)
    ledes: dict[str, str] = {}
    for slug_match in re.finditer(r'"([^"]+)"\s*:\s*\{', case_text):
        slug = slug_match.group(1)
        if "lede" not in case_text[slug_match.start() : slug_match.start() + 800]:
            continue
        obj = slice_balanced(case_text, case_text.find("{", slug_match.start()))
        lede = extract_str(obj, "lede")
        if lede:
            ledes[slug] = lede

    cards = []
    for slug in project_slugs[:6]:
        key = f'"{slug}"'
        idx = meta_block.find(key)
        if idx < 0:
            continue
        obj = slice_balanced(meta_block, meta_block.find("{", idx))
        cards.append(
            {
                "slug": slug,
                "clientName": extract_str(obj, "clientName"),
                "cardTitle": extract_str(obj, "cardTitle"),
                "industry": extract_str(obj, "industry"),
                "technologies": extract_str_list(obj, "technologiesUsed"),
                "summary": ledes.get(slug, ""),
                "status": "Completed",
            }
        )
    return cards


def set_run(run, *, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "Calibri")
    rFonts.set(qn("w:cs"), "Calibri")


def shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="CBD5E1"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", 70), ("left", 90), ("bottom", 70), ("right", 90)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def fill_cell(cell, text: str, *, header=False, size=10, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run(run, size=size, bold=header or bold, color=WHITE if header else NAVY)
    shade_cell(cell, NAVY_HEX if header else "FFFFFF")
    set_cell_borders(cell, NAVY_HEX if header else "CBD5E1")
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    for i, header in enumerate(headers):
        fill_cell(table.rows[0].cells[i], header, header=True, size=10)
        if widths:
            table.rows[0].cells[i].width = Inches(widths[i])
    for r_i, row in enumerate(rows):
        for c_i, value in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            fill_cell(cell, value, size=10)
            if r_i % 2:
                shade_cell(cell, ROW_ALT)
            if widths:
                cell.width = Inches(widths[c_i])
    doc.add_paragraph()


def p(doc, text, *, size=11, bold=False, color=NAVY, space_after=8, italic=False, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        para.alignment = align
    run = para.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return para


def heading(doc, number: str, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(f"{number}  {title}")
    set_run(run, size=16, bold=True, color=NAVY)
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(10)
    run = bar.add_run("━" * 42)
    set_run(run, size=8, color=GOLD)


def label_value(doc, label: str, value: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(f"{label}: ")
    set_run(run, size=11, bold=True, color=NAVY)
    run = para.add_run(value)
    set_run(run, size=11, color=MUTED)


def bullet(doc, text: str, *, bold_prefix: str | None = None):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = para.add_run(f"{bold_prefix} ")
        set_run(run, size=11, bold=True, color=NAVY)
        run = para.add_run(text)
        set_run(run, size=11, color=MUTED)
    else:
        run = para.add_run(text)
        set_run(run, size=11, color=MUTED)


def write_service_page(doc, *, shared: dict, category: dict, subs: list[dict], chapter: int):
    slug = category["slug"]
    label = category["label"]
    page_url = f"{SITE}/{slug}"
    related = [item for item in shared["categories"] if item["slug"] != slug]
    why_us = category["highlights"] or shared["highlights"]
    capabilities_subtitle = (
        f"Focused capabilities within {label.lower()}, each scoped to your timeline and team. "
        "Built by a leading software development company. Open any capability below."
    )
    cta_description = (
        f"Tell us about your {label.lower()} needs, we'll respond within one business day."
    )
    toc = [
        "Header (top bar + navigation)",
        "Hero",
        "Trust numbers",
        "Trusted partners",
        "About / Who we are",
        "What we deliver (capabilities)",
        "Industries we serve",
        "Technology stack",
        "Delivery process",
        "Recent projects",
        "Case studies",
        "Why teams choose us",
        "Team",
        "Explore more capabilities",
        "Testimonials",
        "Blog / insights",
        "FAQs (visible accordion)",
        "Contact",
        "Bottom CTA",
        "Footer + WhatsApp float",
        "Appendix: JSON-LD unique to this page",
    ]

    p(doc, f"CHAPTER {chapter:02d}  ·  {label.upper()}", size=12, bold=True, color=GOLD, space_after=4)
    p(doc, f"{label} service page", size=26, bold=True, space_after=6)
    p(
        doc,
        f"Every visible section on /{slug}, in the same order as the live page. "
        "Includes shared landing modules on this route, plus SEO and JSON-LD copy.",
        size=11,
        color=MUTED,
        space_after=14,
    )

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Live URL", page_url],
            ["Meta title", category["pageTitle"]],
            ["Meta description", category["metaDescription"]],
            ["Canonical", page_url],
            ["Open Graph title", category["pageTitle"]],
            ["Open Graph description", category["metaDescription"]],
            ["Hero image", category["heroImage"]],
            ["Hero image alt", category["heroAlt"]],
        ],
        [1.8, 5.0],
    )

    p(doc, "Page order", size=13, bold=True, space_after=6)
    for i, item in enumerate(toc, 1):
        bullet(doc, f"{i:02d}.  {item}")

    heading(doc, "01", "Header")
    p(doc, "Top bar", size=13, bold=True, space_after=4)
    bullet(doc, "Islamabad, Pakistan")
    bullet(doc, "+92 371 0510083")
    bullet(doc, "info@nextsoftwaredevelopment.com")
    p(
        doc,
        "Social icons: LinkedIn, Facebook, YouTube, X (Twitter), GitHub, Clutch.",
        color=MUTED,
    )

    p(doc, "Primary navigation", size=13, bold=True, space_after=4)
    service_children = ", ".join(item["label"] for item in shared["categories"])
    nav_rows = [
        ["Services", "/software-development", service_children],
        ["Locations", "/location", "Pakistan"],
        ["Industries", "/industries", ""],
        ["Solutions", "/solutions", ""],
        ["Portfolio", "/projects", "Case Studies, Recent Projects"],
        ["About", "/about", "Our Team, Gallery"],
        ["Resources", "/blog", "Blog, FAQ"],
        ["Contact", "/contact", ""],
    ]
    add_table(doc, ["Item", "URL", "Children"], nav_rows, [1.6, 2.2, 3.0])
    label_value(doc, "Header CTA", "Get a Free Quote → /contact")

    heading(doc, "02", "Hero")
    label_value(doc, "Breadcrumb", f"Home  /  {label}")
    label_value(doc, "Eyebrow", label)
    label_value(doc, "H1", category["tagline"])
    p(doc, category["description"], space_after=10)
    label_value(doc, "Primary CTA", "Get a Free Quote → /contact")
    label_value(doc, "Secondary CTA", "See Our Work → /projects")
    label_value(doc, "Hero image", f'{category["heroImage"]}  ({category["heroAlt"]})')
    p(
        doc,
        "Rating badges (Facebook, Clutch, Trustpilot, Google) sit beside the hero image.",
        color=MUTED,
    )

    heading(doc, "03", "Trust numbers")
    add_table(
        doc,
        ["Number", "Label", "Detail"],
        [[item["value"], item["label"], item["detail"]] for item in shared["stats"]],
        [1.2, 1.8, 4.8],
    )

    heading(doc, "04", "Trusted partners")
    p(doc, "Logo marquee. Accessible list of partner names:", color=MUTED)
    p(doc, ", ".join(item["name"] for item in shared["clients"] if item["name"]), space_after=10)

    heading(doc, "05", "About / Who we are")
    label_value(doc, "Overline", "About us")
    label_value(doc, "H2", "Who we are?")
    for paragraph in shared["about"]["paragraphs"]:
        p(doc, paragraph, space_after=8)
    label_value(doc, "CTA", f'{shared["about"]["teamCta"]} → /team')
    label_value(doc, "Image alt", shared["about"]["imageAlt"])
    p(doc, "Value cards", size=13, bold=True, space_after=6)
    add_table(
        doc,
        ["Title", "Description"],
        [[item["title"], item["description"]] for item in shared["about"]["values"]],
        [2.2, 5.6],
    )

    heading(doc, "06", "What we deliver (capabilities)")
    label_value(doc, "H2", "What we deliver")
    p(doc, capabilities_subtitle, space_after=10)
    for item in subs:
        p(doc, item["label"], size=13, bold=True, space_after=2)
        label_value(doc, "Card copy", item["description"])
        label_value(doc, "URL", f"{SITE}/{slug}/{item['slug']}")
        label_value(doc, "Card CTA", "Learn more")
        label_value(doc, "Eyebrow on card", "Capability")

    heading(doc, "07", "Industries we serve")
    label_value(doc, "Overline", "Industries we serve")
    label_value(doc, "H2", "Software solutions for every sector")
    p(
        doc,
        "As a best software house for complex industry work, we tailor products to the workflows, "
        "compliance needs, and growth goals of every sector we support.",
        space_after=10,
    )
    add_table(
        doc,
        ["Industry", "Description", "Learn more URL"],
        [
            [item["industry"], item["description"], f"{SITE}/industries/{item['slug']}"]
            for item in shared["industries"]
        ],
        [1.7, 3.6, 2.5],
    )
    for item in shared["industries"]:
        label_value(doc, f"CTA for {item['industry']}", f"Learn About {item['industry']}")

    heading(doc, "08", "Technology stack")
    label_value(doc, "Overline", "Technology stack")
    label_value(doc, "H2", "Built with proven, modern tools")
    label_value(doc, "Badge", "Technology built around your Business")
    p(doc, shared["tech_intro"], space_after=10)
    p(doc, "Interactive note on the stack explorer: Drag · Explore", italic=True, color=MUTED)
    for index, group in enumerate(shared["tech_groups"], 1):
        p(doc, f"{index:02d}.  {group['label']}", size=13, bold=True, space_after=3)
        p(doc, group["description"], space_after=4)
        p(doc, "Technologies: " + ", ".join(group["logos"]), color=MUTED, space_after=10)

    heading(doc, "09", "Delivery process")
    label_value(doc, "Overline", "How we work")
    label_value(doc, "H2", "Our delivery process transparent, agile, and built around you")
    p(
        doc,
        "Many teams talk about best practices. As a top rated software house, we show you exactly what delivery looks like at every stage of your project.",
        space_after=10,
    )
    add_table(
        doc,
        ["Step", "Title", "Description"],
        [[item["step"], item["title"], item["description"]] for item in shared["process_steps"]],
        [0.8, 2.4, 4.6],
    )
    label_value(doc, "CTA", "Start your project → /contact")

    heading(doc, "10", "Recent projects")
    label_value(doc, "Overline", "Software house portfolio")
    label_value(doc, "H2", "Recent projects")
    p(
        doc,
        "Carousel of all showcase projects. Each project includes title, category, description, highlights, slides, and a View case study link.",
        color=MUTED,
    )
    for index, item in enumerate(shared["projects"], 1):
        p(doc, f"{index:02d}.  {item['title']}", size=13, bold=True, space_after=2)
        label_value(doc, "Category", item["category"])
        p(doc, item["description"], space_after=4)
        highlights_list = shared["project_highlights"].get(item["slug"], [])
        if highlights_list:
            p(doc, "Highlights", size=11, bold=True, space_after=2)
            for point in highlights_list:
                bullet(doc, point)
        slides = shared["project_slides"].get(item["slug"], [])
        if slides:
            p(doc, "Slides on this project", size=11, bold=True, space_after=2)
            for slide in slides:
                bullet(
                    doc,
                    f"{slide['caption']}  ({slide['src']})" if slide.get("src") else slide["caption"],
                    bold_prefix=f"{slide['label']}:",
                )
        label_value(doc, "URL", f"{SITE}/projects/{item['slug']}")
        label_value(doc, "CTA", "View case study")
    label_value(doc, "Browse label", "Browse projects")
    label_value(doc, "Section CTA", "View all case studies and projects → /projects")

    heading(doc, "11", "Case studies")
    label_value(doc, "Overline", "Case studies")
    label_value(doc, "H2", "Results from real engagements")
    p(
        doc,
        "Selected case studies showing how we deliver outcomes across products and industries.",
        space_after=10,
    )
    p(doc, "First 6 case studies are shown on this page.", italic=True, color=MUTED)
    for item in shared["case_cards"]:
        p(doc, item["cardTitle"], size=13, bold=True, space_after=2)
        label_value(doc, "Client", item["clientName"])
        label_value(doc, "Category", item["industry"])
        label_value(doc, "Status", item["status"])
        if item["summary"]:
            p(doc, item["summary"], space_after=4)
        if item["technologies"]:
            label_value(doc, "Technologies", ", ".join(item["technologies"]))
        label_value(doc, "URL", f"{SITE}/case-study/{item['slug']}")
        label_value(doc, "CTA", "Read case study")

    heading(doc, "12", "Why teams choose us")
    label_value(doc, "H2", "Why teams choose us")
    add_table(
        doc,
        ["Title", "Description"],
        [[item["title"], item["description"]] for item in why_us],
        [2.4, 5.4],
    )

    heading(doc, "13", "Team")
    label_value(doc, "Overline", "Leading software company team")
    label_value(doc, "H2", "A small team.")
    p(doc, shared["team_intro"], space_after=10)
    add_table(
        doc,
        ["Name", "Role", "Bio"],
        [[item["name"], item["role"], item["bio"]] for item in shared["team_members"]],
        [1.8, 2.0, 4.0],
    )

    heading(doc, "14", "Explore more capabilities")
    label_value(doc, "H2", "Explore more capabilities")
    add_table(
        doc,
        ["Service", "Tagline", "URL"],
        [[item["label"], item["tagline"], f"{SITE}/{item['slug']}"] for item in related],
        [2.2, 3.4, 2.2],
    )

    heading(doc, "15", "Testimonials")
    label_value(doc, "Overline", "Client feedback")
    label_value(doc, "H2", "What partners say")
    p(doc, "Cards rotate three at a time on the live page. Full set is included below.", italic=True, color=MUTED)
    for item in shared["testimonials"]:
        p(doc, f"“{item['quote']}”", italic=True, space_after=3)
        p(doc, f"{item['author']}  ·  {item['role']}", size=10, bold=True, color=MUTED, space_after=10)

    heading(doc, "16", "Blog / insights")
    label_value(doc, "Overline", "Insights from our software house")
    label_value(doc, "H2", "Insights from our team")
    p(
        doc,
        "Practical notes on product delivery, discovery, and building systems that scale, written by the people behind the work.",
        space_after=10,
    )
    for item in shared["blogs"]:
        p(doc, item["title"], size=13, bold=True, space_after=2)
        label_value(doc, "Category", item["category"])
        label_value(doc, "Date / time", f"{item['date']}  ·  {item['readTime']}")
        p(doc, item["excerpt"], space_after=4)
        label_value(doc, "URL", f"{SITE}/blog/{item['slug']}")

    heading(doc, "17", "FAQs (visible accordion)")
    label_value(doc, "Overline", "Software development company FAQs")
    label_value(doc, "H2", "Questions, answered")
    p(doc, shared["faq_intro"], space_after=10)
    if category["faqs"]:
        p(
            doc,
            "This is the accordion visitors see. A separate category FAQ set is used in JSON-LD and is listed in the appendix.",
            italic=True,
            color=MUTED,
        )
    for index, item in enumerate(shared["faqs"], 1):
        p(doc, f"{index:02d}.  {item['question']}", size=12, bold=True, space_after=2)
        label_value(doc, "Tag", item["tag"])
        p(doc, item["answer"], space_after=10)

    heading(doc, "18", "Contact")
    contact = shared["contact"]
    label_value(doc, "Overline", contact["overline"])
    label_value(doc, "H2", f"{contact['titleBefore']} {contact['titleEmphasis']}")
    p(doc, contact["subtext"], space_after=8)
    p(doc, "Short form fields on this page", size=13, bold=True, space_after=4)
    bullet(doc, "Name")
    bullet(doc, "Email")
    label_value(doc, "Submit", contact["cta"])
    p(doc, contact["reassurance"], italic=True, color=MUTED)
    p(
        doc,
        "The short form saves name/email and continues to /contact. The full contact page also includes:",
        space_after=6,
    )
    p(doc, "Project type options", size=12, bold=True, space_after=3)
    for item in shared["project_types"]:
        bullet(doc, item["label"])
    p(doc, "Budget options", size=12, bold=True, space_after=3)
    for item in shared["budgets"]:
        bullet(doc, item["label"])
    p(doc, "Timeline options", size=12, bold=True, space_after=3)
    for item in shared["timelines"]:
        bullet(doc, item["label"])

    heading(doc, "19", "Bottom CTA")
    label_value(doc, "H2", "Ready to build with Software Development Company?")
    p(doc, cta_description, space_after=6)
    label_value(doc, "Button", "Get in touch → /contact")

    heading(doc, "20", "Footer and floating actions")
    p(doc, shared["footer_brand"], space_after=8)
    bullet(doc, "Email: info@nextsoftwaredevelopment.com")
    bullet(doc, "Phone: +92 371 0510083")
    bullet(doc, "Location: Islamabad, Pakistan")
    p(doc, "Footer columns", size=13, bold=True, space_after=4)
    p(doc, "Services — mobile accordion shows the first 6; desktop shows the full list.", color=MUTED)
    p(doc, "Mobile (first 6)", size=12, bold=True, space_after=3)
    for item in shared["categories"][:6]:
        bullet(doc, f"{item['label']}  →  {SITE}/{item['slug']}")
    p(doc, "Desktop (all services)", size=12, bold=True, space_after=3)
    for item in shared["categories"]:
        bullet(doc, f"{item['label']}  →  {SITE}/{item['slug']}")
    p(doc, "Company", size=12, bold=True, space_after=3)
    for company_label, href in [
        ("About Us", "/about"),
        ("Our Work", "/projects"),
        ("Industries", "/industries"),
        ("Solutions", "/solutions"),
        ("Blog", "/blog"),
        ("FAQ", "/faqs"),
        ("Contact", "/contact"),
    ]:
        bullet(doc, f"{company_label}  →  {SITE}{href}")
    p(doc, "Connect", size=12, bold=True, space_after=3)
    for connect_label, href in [
        ("LinkedIn", "https://www.linkedin.com/company/nextsoftwaredevelopmentcompany"),
        ("Facebook", "https://www.facebook.com/nextsoftwaredevelopmentcompany"),
        ("YouTube", "https://www.youtube.com/@nextsoftwaredevelopmentcompany"),
        ("Twitter / X", "https://x.com/NextSoftwaree"),
        ("GitHub", "https://github.com/nextsoftwaredevelopmentcompany"),
        ("Clutch", "https://clutch.co/profile/next-software-development-company"),
    ]:
        bullet(doc, f"{connect_label}  →  {href}")
    label_value(doc, "Legal", "Privacy Policy → /privacy   ·   Terms of Service → /terms")
    label_value(doc, "Copyright", "© 2026 Next Software Development Company. All rights reserved.")
    p(doc, "WhatsApp float", size=13, bold=True, space_after=4)
    bullet(doc, "Label: WhatsApp Us at +92 371 0510083")
    bullet(doc, "Prefilled message: Hi, I'd like to discuss a software project.")
    bullet(doc, "Link: https://wa.me/923710510083")

    heading(doc, "21", "Appendix — JSON-LD unique to this page")
    if category["faqs"]:
        p(
            doc,
            f"These questions are defined on the {label} category and output as FAQPage structured data. "
            "They are not the same as the visible FAQ accordion above.",
            italic=True,
            color=MUTED,
            space_after=10,
        )
        for index, item in enumerate(category["faqs"], 1):
            p(doc, f"{index:02d}.  {item['question']}", size=12, bold=True, space_after=2)
            label_value(doc, "Tag", item["tag"])
            p(doc, item["answer"], space_after=10)
    else:
        p(
            doc,
            f"{label} does not define category-specific FAQPage JSON-LD. The visible accordion still uses the shared FAQ set above.",
            italic=True,
            color=MUTED,
            space_after=10,
        )

    p(doc, "Also in JSON-LD on this URL", size=13, bold=True, space_after=4)
    if slug == "software-development":
        bullet(doc, "ProfessionalService schema (software-development only)")
    bullet(doc, f'Service schema — name: {label}')
    bullet(doc, f"BreadcrumbList: Home → {label}")
    if category["faqs"]:
        bullet(doc, f"FAQPage using the {label} questions listed above")

    p(
        doc,
        f"{label} chapter generated from live source files. {len(subs)} capabilities, "
        f"{len(shared['industries'])} industries, {len(shared['projects'])} projects, "
        f"{len(shared['case_cards'])} case-study cards, {len(shared['testimonials'])} testimonials, "
        f"{len(shared['faqs'])} visible FAQs, {len(category['faqs'])} JSON-LD FAQs.",
        size=9,
        color=MUTED,
        space_after=0,
    )


def main():
    services_text = (ROOT / "data" / "services.tsx").read_text(encoding="utf-8")
    sub_text = (ROOT / "data" / "subServices.tsx").read_text(encoding="utf-8")
    home_text = (ROOT / "data" / "homepage-content.ts").read_text(encoding="utf-8")
    industries_text = (ROOT / "data" / "industriesPage.ts").read_text(encoding="utf-8")
    projects_text = (ROOT / "data" / "projects.ts").read_text(encoding="utf-8")
    case_text = (ROOT / "data" / "caseStudy.ts").read_text(encoding="utf-8")
    navbar_text = (ROOT / "components" / "landing" / "navbar.tsx").read_text(encoding="utf-8")

    categories = parse_define_categories(services_text)
    default_highlights = parse_highlights(services_text)
    about = parse_about(home_text)
    stats = parse_named_object_array(home_text, "companyStats", ["value", "label", "detail"])
    clients = parse_named_object_array(home_text, "clients", ["name"])
    process_steps = parse_named_object_array(
        home_text, "processSteps", ["step", "shortLabel", "title", "description"]
    )
    tech_groups = parse_tech_groups(home_text)
    team_members = parse_named_object_array(home_text, "teamMembers", ["name", "role", "bio"])
    testimonials = parse_named_object_array(home_text, "testimonials", ["quote", "author", "role"])
    faqs = parse_named_object_array(home_text, "faqs", ["tag", "question", "answer"])
    faq_intro = parse_export_string(home_text, "faqIntro")
    team_intro = parse_export_string(home_text, "teamIntro")
    tech_intro = parse_export_string(home_text, "techStackIntro")
    contact = parse_home_contact(home_text)
    industries = parse_named_object_array(
        industries_text, "industries", ["industry", "description", "slug"]
    )
    projects = parse_named_object_array(
        home_text, "projects", ["slug", "title", "category", "description"]
    )
    project_highlights = parse_project_highlights(projects_text)
    project_slides = parse_project_slides(projects_text)
    case_cards = parse_case_study_cards(case_text, [item["slug"] for item in projects])
    blogs = parse_named_object_array(
        home_text, "blogPosts", ["slug", "title", "excerpt", "date", "readTime", "category"]
    )[:3]
    project_types = parse_named_object_array(home_text, "contactProjectTypes", ["label"])
    budgets = parse_named_object_array(home_text, "contactBudgetRanges", ["label"])
    timelines = parse_named_object_array(home_text, "contactTimelineOptions", ["label"])
    footer_brand = extract_str(
        slice_balanced(
            navbar_text,
            navbar_text.find("export const footerBrand") + len("export const footerBrand = {") - 1,
        )
        if "export const footerBrand" in navbar_text
        else "",
        "description",
    )
    if not footer_brand:
        footer_brand = extract_str(navbar_text[navbar_text.find("footerBrand") :], "description")

    shared = {
        "categories": categories,
        "highlights": default_highlights,
        "about": about,
        "stats": stats,
        "clients": clients,
        "process_steps": process_steps,
        "tech_groups": tech_groups,
        "team_members": team_members,
        "testimonials": testimonials,
        "faqs": faqs,
        "faq_intro": faq_intro,
        "team_intro": team_intro,
        "tech_intro": tech_intro,
        "contact": contact,
        "industries": industries,
        "projects": projects,
        "project_highlights": project_highlights,
        "project_slides": project_slides,
        "case_cards": case_cards,
        "blogs": blogs,
        "project_types": project_types,
        "budgets": budgets,
        "timelines": timelines,
        "footer_brand": footer_brand
        or "Next Software Development Company helps teams across the USA, UK, UAE, Canada, Australia, and beyond ship reliable digital products with senior engineering and clear communication.",
    }

    targets = []
    for slug in TARGET_SLUGS:
        category = next((cat for cat in categories if cat["slug"] == slug), None)
        if not category:
            raise SystemExit(f"Missing category: {slug}")
        subs = parse_sub_services(sub_text, slug)
        targets.append((category, subs))

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(
        "Next Software Development Company  ·  Software, Mobile & Web Development pages"
    )
    set_run(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Full on-page content  ·  nextsoftwaredevelopment.com")
    set_run(run, size=8, color=MUTED)

    p(doc, "MAIN SERVICE PAGES", size=12, bold=True, color=GOLD, space_after=4)
    p(doc, "Software, Mobile & Web Development", size=26, bold=True, space_after=6)
    p(
        doc,
        "Complete on-page content for three live service routes. Each chapter is a full dump "
        "of that page in the same section order visitors see, including shared modules, SEO, and JSON-LD.",
        size=11,
        color=MUTED,
        space_after=14,
    )
    add_table(
        doc,
        ["Chapter", "Service", "Live URL"],
        [
            [f"{index:02d}", category["label"], f"{SITE}/{category['slug']}"]
            for index, (category, _) in enumerate(targets, 1)
        ],
        [1.1, 2.8, 3.9],
    )
    p(doc, "How to use this file", size=13, bold=True, space_after=4)
    bullet(doc, "Chapter 01 is the full Software Development page.")
    bullet(doc, "Chapter 02 is the full Mobile App Development page.")
    bullet(doc, "Chapter 03 is the full Web Development page.")
    bullet(doc, "Shared chrome (header, trust, about, industries, stack, process, projects, team, FAQs, contact, footer) is included in every chapter so nothing is missing from any page.")

    missing = []
    if len(categories) != 16:
        missing.append(f"expected 16 main services, got {len(categories)}")
    if len(projects) < 10:
        missing.append(f"expected 12 projects, got {len(projects)}")
    for category, subs in targets:
        if len(subs) != 6:
            missing.append(f"{category['slug']}: expected 6 capabilities, got {len(subs)}")
    if missing:
        raise SystemExit("Extraction incomplete: " + "; ".join(missing))

    for index, (category, subs) in enumerate(targets, 1):
        doc.add_page_break()
        write_service_page(doc, shared=shared, category=category, subs=subs, chapter=index)

    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")
    for category, subs in targets:
        print(
            f"- {category['label']}: {len(subs)} capabilities, "
            f"{len(category['faqs'])} JSON-LD FAQs"
        )
    print(f"Industries: {len(industries)}")
    print(f"Projects: {len(projects)}")
    print(f"Project slides: {sum(len(v) for v in project_slides.values())}")
    print(f"Case studies: {len(case_cards)}")
    print(f"Testimonials: {len(testimonials)}")
    print(f"Visible FAQs: {len(faqs)}")
    print(f"Team: {len(team_members)}")


if __name__ == "__main__":
    main()
